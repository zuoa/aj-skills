#!/usr/bin/env python3
"""Convert a Markdown copyright document to a consistently styled DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path


BODY_FONT_EAST_ASIA = "宋体"
BODY_FONT_LATIN = "Times New Roman"
HEADING_FONT_EAST_ASIA = "黑体"
HEADING_FONT_LATIN = "Arial"
MONO_FONT = "Consolas"
TABLE_WIDTH_DXA = 9026  # A4 width minus 2.54 cm margins.
TABLE_INDENT_DXA = 100

# Base preset: standard_business_brief.
# Named overrides for Chinese software-registration documents:
# - cn_registration_a4: A4 page and 9026 DXA usable width.
# - cn_technical_typography: explicit SimSun/SimHei East Asian fonts, 10.5 pt body.
# - cn_table_alignment: 100 DXA start/end cell margins and matching table indent.


def import_docx():
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Inches, Pt, RGBColor
    except ImportError as exc:
        raise SystemExit("Missing dependency: pip install python-docx") from exc
    return (
        Document,
        Cm,
        Inches,
        Pt,
        RGBColor,
        WD_ALIGN_PARAGRAPH,
        WD_CELL_VERTICAL_ALIGNMENT,
        WD_TABLE_ALIGNMENT,
        OxmlElement,
        qn,
    )


def resolve_template(template: str | None, cwd: Path, markdown_path: Path | None = None) -> Path | None:
    candidates: list[Path] = []
    if template:
        candidates.append(Path(template).expanduser())
    manual_names = ["操作手册模版.docx", "操作手册模板.docx"]
    design_names = ["软件设计说明书模版.docx", "软件设计说明书模板.docx"]
    if markdown_path and "软件设计说明书" in markdown_path.name:
        names = design_names + manual_names
    else:
        names = manual_names + design_names
    roots = [
        cwd / "reference",
        cwd / "refence",
        cwd / "refrence",
        Path.home() / "aj-skills" / "reference",
        Path.home() / "aj-skills" / "refence",
        Path.home() / "aj-skills" / "refrence",
    ]
    for root in roots:
        for name in names:
            candidates.append(root / name)
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def is_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def set_rfonts(rpr, east_asia: str, latin: str) -> None:
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        qname = qn(f"w:{attr}")
        if qname in rfonts.attrib:
            del rfonts.attrib[qname]
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)


def set_run_font(run, east_asia: str, latin: str, size: float | None = None) -> None:
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    set_rfonts(run._element.get_or_add_rPr(), east_asia, latin)


def configure_style(style, east_asia: str, latin: str, size: float, bold: bool = False) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style._element.get_or_add_rPr()
    set_rfonts(rpr, east_asia, latin)
    for tag in ("w:sz", "w:szCs"):
        node = rpr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            rpr.append(node)
        node.set(qn("w:val"), str(int(size * 2)))
    for tag in ("w:spacing", "w:kern"):
        node = rpr.find(qn(tag))
        if node is not None:
            rpr.remove(node)


def configure_document_styles(document) -> None:
    """Apply an A4 Chinese technical-document typography preset."""
    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    normal = document.styles["Normal"]
    configure_style(normal, BODY_FONT_EAST_ASIA, BODY_FONT_LATIN, 10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.25

    title = document.styles["Title"]
    configure_style(title, HEADING_FONT_EAST_ASIA, HEADING_FONT_LATIN, 20, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(18)
    title.paragraph_format.line_spacing = 1.15
    title.paragraph_format.keep_with_next = True
    title_ppr = title._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)
    contextual_spacing = title_ppr.find(qn("w:contextualSpacing"))
    if contextual_spacing is not None:
        title_ppr.remove(contextual_spacing)

    heading_tokens = {
        "Heading 1": (16, 16, 8),
        "Heading 2": (14, 12, 6),
        "Heading 3": (12, 10, 5),
        "Heading 4": (11, 8, 4),
    }
    for name, (size, before, after) in heading_tokens.items():
        style = document.styles[name]
        configure_style(style, HEADING_FONT_EAST_ASIA, HEADING_FONT_LATIN, size, bold=True)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        configure_style(style, BODY_FONT_EAST_ASIA, BODY_FONT_LATIN, 10.5)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.25

    if "Caption" in document.styles:
        caption = document.styles["Caption"]
        configure_style(caption, BODY_FONT_EAST_ASIA, BODY_FONT_LATIN, 9)
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.first_line_indent = Pt(0)
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(6)


def text_display_width(text: str) -> int:
    return sum(2 if "\u2e80" <= char <= "\u9fff" else 1 for char in text)


def calculate_column_widths(rows: list[list[str]], column_count: int) -> list[int]:
    weights = []
    for column_index in range(column_count):
        widest = max(
            (text_display_width(row[column_index]) for row in rows if column_index < len(row)),
            default=4,
        )
        weights.append(max(4, min(widest, 40)))
    minimum = max(500, min(900, TABLE_WIDTH_DXA // max(column_count * 2, 1)))
    remaining = TABLE_WIDTH_DXA - minimum * column_count
    if remaining <= 0:
        widths = [TABLE_WIDTH_DXA // column_count] * column_count
    else:
        weight_total = sum(weights)
        widths = [minimum + int(remaining * weight / weight_total) for weight in weights]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "B7C1CC")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, margin in (("top", 80), ("start", 100), ("bottom", 80), ("end", 100)):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(margin))
                node.set(qn("w:type"), "dxa")


def add_table(document, lines: list[str]) -> None:
    rows = [split_table_row(line) for line in lines]
    rows = [row for row in rows if not is_separator_row(row)]
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=cols)
    table.style = "Table Grid"
    for row_index, row_cells in enumerate(rows):
        row = table.add_row().cells
        for idx in range(cols):
            if idx < len(row_cells):
                add_inline_markdown(row[idx].paragraphs[0], row_cells[idx])
            paragraph = row[idx].paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            for run in paragraph.runs:
                latin_font = MONO_FONT if run.font.name == MONO_FONT else BODY_FONT_LATIN
                set_run_font(run, BODY_FONT_EAST_ASIA, latin_font, 9.5)
                if row_index == 0:
                    run.bold = True
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F4F7")
                row[idx]._tc.get_or_add_tcPr().append(shading)
    if table.rows:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        tr_pr.append(repeat)
    set_table_geometry(table, calculate_column_widths(rows, cols))


def add_image(document, markdown_image: re.Match[str], base_dir: Path, Inches) -> bool:
    alt = markdown_image.group(1).strip()
    raw_path = markdown_image.group(2).strip().strip("<>").strip('"').strip("'")
    image_path = Path(raw_path)
    if not image_path.is_absolute():
        image_path = (base_dir / image_path).resolve()
    if not image_path.exists():
        document.add_paragraph(f"[图片缺失: {alt} - {raw_path}]")
        return False
    document.add_picture(str(image_path), width=Inches(6.1))
    if alt:
        caption = document.add_paragraph(style="Caption")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(alt)
        set_run_font(run, BODY_FONT_EAST_ASIA, BODY_FONT_LATIN, 9)
    return True


def add_code_block(document, code_lines: list[str]) -> None:
    for line in code_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(0)
        run = paragraph.add_run(line if line else " ")
        set_run_font(run, BODY_FONT_EAST_ASIA, MONO_FONT, 8.5)


def add_inline_markdown(paragraph, text: str) -> None:
    """Render a small subset of inline Markdown into docx runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, BODY_FONT_EAST_ASIA, MONO_FONT)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_paragraph_with_inline(document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style) if style else document.add_paragraph()
    add_inline_markdown(paragraph, text)
    return paragraph


def convert_markdown(markdown_path: Path, output_path: Path, template_path: Path | None, auto_lists: bool = True) -> None:
    imports = import_docx()
    global Cm, Inches, Pt, RGBColor, WD_ALIGN_PARAGRAPH, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, OxmlElement, qn
    (
        Document,
        Cm,
        Inches,
        Pt,
        RGBColor,
        WD_ALIGN_PARAGRAPH,
        WD_CELL_VERTICAL_ALIGNMENT,
        WD_TABLE_ALIGNMENT,
        OxmlElement,
        qn,
    ) = imports

    document = Document(str(template_path)) if template_path else Document()
    configure_document_styles(document)
    if template_path and any(p.text.strip() for p in document.paragraphs):
        document.add_page_break()

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buffer: list[str] = []
    title_written = False
    image_re = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_buffer)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if is_table_line(line):
            table_lines = []
            while i < len(lines) and is_table_line(lines[i]):
                table_lines.append(lines[i])
                i += 1
            add_table(document, table_lines)
            continue

        image_match = image_re.fullmatch(stripped)
        if image_match:
            add_image(document, image_match, markdown_path.parent, Inches)
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)), 4)
            heading_text = heading.group(2).strip()
            if level == 1 and not title_written and ("说明书" in heading_text or "操作手册" in heading_text):
                paragraph = document.add_paragraph(style="Title")
                add_inline_markdown(paragraph, heading_text)
                title_written = True
            else:
                document.add_heading(heading_text, level=level)
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            if auto_lists:
                add_paragraph_with_inline(document, bullet.group(1), style="List Bullet")
            else:
                add_paragraph_with_inline(document, f"• {bullet.group(1)}")
            i += 1
            continue

        ordered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if ordered:
            if auto_lists:
                add_paragraph_with_inline(document, ordered.group(1), style="List Number")
            else:
                add_paragraph_with_inline(document, stripped)
            i += 1
            continue

        add_paragraph_with_inline(document, stripped)
        i += 1

    if code_buffer:
        add_code_block(document, code_buffer)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--template")
    parser.add_argument("--auto-lists", action="store_true", help=argparse.SUPPRESS)
    parser.add_argument(
        "--plain-list-markers",
        action="store_true",
        help="Render list markers as plain text for compatibility. Word list styles are used by default.",
    )
    args = parser.parse_args()

    markdown_path = args.input
    template_path = resolve_template(args.template, Path.cwd(), markdown_path)
    if template_path:
        print(f"Using template: {template_path}")
    else:
        print("Template not found; generating docx without template.")
    convert_markdown(markdown_path, args.output, template_path, auto_lists=not args.plain_list_markers)
    print(f"Wrote {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
