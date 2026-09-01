#!/usr/bin/env python3
"""Build a source-code docx from numbered 05.code/*.txt files."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path


SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from validate_outputs import ValidationError, validate_code_files, validate_file_exists, validate_originality_report


def import_docx():
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise SystemExit("Missing dependency: pip install python-docx") from exc
    return Document, WD_STYLE_TYPE, WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, OxmlElement, qn, Cm, Pt


def resolve_template(template: str | None, cwd: Path) -> Path | None:
    if template:
        path = Path(template).expanduser()
        if path.exists():
            return path
    names = ["代码文档模版.docx", "代码文档模板.docx", "操作手册模版.docx", "操作手册模板.docx"]
    roots = [
        cwd / "reference",
        cwd / "refence",
        cwd / "refrence",
        Path.home() / "aj-skills" / "reference",
        Path.home() / "aj-skills" / "refence",
        Path.home() / "aj-skills" / "refrence",
    ]
    for root in roots:
        for name in names:
            candidate = root / name
            if candidate.exists():
                return candidate
    return None


def set_run_font(run, qn, pt, font_name: str = "宋体", font_size: float = 9) -> None:
    run.font.name = font_name
    run.font.size = pt(font_size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def ensure_code_style(document, style_type, align, qn, pt):
    styles = document.styles
    try:
        style = styles["AJ Code"]
    except KeyError:
        style = styles.add_style("AJ Code", style_type.PARAGRAPH)
    style.font.name = "宋体"
    style.font.size = pt(9)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.paragraph_format.alignment = align.LEFT
    style.paragraph_format.line_spacing = 1
    style.paragraph_format.space_before = pt(0)
    style.paragraph_format.space_after = pt(0)
    return style


def clear_body(document, qn) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_page(document, cm) -> None:
    section = document.sections[0]
    section.page_width = cm(21)
    section.page_height = cm(29.7)
    section.top_margin = cm(1.8)
    section.bottom_margin = cm(1.6)
    section.left_margin = cm(2.0)
    section.right_margin = cm(2.0)
    section.header_distance = cm(0.8)
    section.footer_distance = cm(0.8)


def clear_paragraph(paragraph) -> None:
    paragraph._element.clear_content()


def add_page_number(paragraph, oxml_element, qn, pt) -> None:
    run = paragraph.add_run("第 ")
    set_run_font(run, qn, pt)

    begin = oxml_element("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = oxml_element("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = oxml_element("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = oxml_element("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    field_run = paragraph.add_run()
    field_run._r.append(begin)
    field_run._r.append(instr)
    field_run._r.append(separate)
    field_run._r.append(end)

    run = paragraph.add_run(" 页")
    set_run_font(run, qn, pt)


def configure_header(document, software_name: str, software_version: str, align, tab_align, oxml_element, qn, pt) -> None:
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    clear_paragraph(paragraph)
    paragraph.paragraph_format.alignment = align.LEFT
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_before = pt(0)
    paragraph.paragraph_format.space_after = pt(0)
    content_width = section.page_width - section.left_margin - section.right_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(content_width, tab_align.RIGHT)

    title_run = paragraph.add_run(f"{software_name} {software_version}")
    set_run_font(title_run, qn, pt)
    tab_run = paragraph.add_run("\t")
    set_run_font(tab_run, qn, pt)
    add_page_number(paragraph, oxml_element, qn, pt)


def code_files(code_dir: Path, source_manifest: Path | None = None) -> list[Path]:
    return validate_code_files(code_dir, source_manifest=source_manifest)


def infer_software_name(output_path: Path) -> str:
    stem = output_path.stem
    if stem.endswith("_代码"):
        return stem[: -len("_代码")]
    return stem


def source_lines_for_deposit(paths: list[Path], required_lines: int) -> list[str]:
    lines: list[str] = []
    errors: list[str] = []
    for path in paths:
        for line_number, raw_line in enumerate(path.read_text(encoding="utf-8").splitlines(), 1):
            if "copyright" in raw_line.lower():
                errors.append(f"{path.name}:{line_number}: contains copyright")
            if not raw_line.strip():
                continue
            lines.append(raw_line.rstrip())

    if errors:
        preview = "; ".join(errors[:20])
        if len(errors) > 20:
            preview += f"; ... and {len(errors) - 20} more"
        raise ValidationError("source code content: " + preview)
    if len(lines) < required_lines:
        raise ValidationError(f"source code content: expected at least {required_lines} non-empty source lines, found {len(lines)}")
    if len(lines) == required_lines:
        return lines

    first_half = required_lines // 2
    last_half = required_lines - first_half
    return lines[:first_half] + lines[-last_half:]


def add_deposit_code_pages(document, lines: list[str], style_name: str, line_numbers: bool, lines_per_page: int, wd_break, qn, pt) -> None:
    for index, line in enumerate(lines, start=1):
        paragraph = document.add_paragraph(style=style_name)
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.first_line_indent = None
        if line_numbers:
            number_run = paragraph.add_run(f"{index:04d}  ")
            set_run_font(number_run, qn, pt)
        run = paragraph.add_run(line)
        set_run_font(run, qn, pt)
        if index % lines_per_page == 0 and index != len(lines):
            run.add_break(wd_break.PAGE)


def build_docx(
    code_dir: Path,
    output_path: Path,
    template_path: Path | None,
    line_numbers: bool,
    software_name: str | None,
    software_version: str,
    source_manifest: Path | None = None,
    originality_report: Path | None = None,
) -> None:
    if source_manifest is None or originality_report is None:
        raise ValidationError("code docx requires a source manifest and a passed originality report")
    validate_originality_report(originality_report, code_dir, source_manifest)
    Document, style_type, align, wd_break, tab_align, oxml_element, qn, cm, pt = import_docx()
    document = Document(str(template_path)) if template_path else Document()
    clear_body(document, qn)
    configure_page(document, cm)
    ensure_code_style(document, style_type, align, qn, pt)
    header_name = software_name or infer_software_name(output_path)
    configure_header(document, header_name, software_version, align, tab_align, oxml_element, qn, pt)

    lines_per_page = 50
    total_pages = 60
    selected_lines = source_lines_for_deposit(code_files(code_dir, source_manifest), lines_per_page * total_pages)
    add_deposit_code_pages(document, selected_lines, "AJ Code", line_numbers, lines_per_page, wd_break, qn, pt)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    validate_file_exists(output_path, "code docx")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--code-dir", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--template")
    parser.add_argument("--software-name")
    parser.add_argument("--software-version", default="V1.0")
    parser.add_argument("--source-manifest", required=True, type=Path)
    parser.add_argument("--originality-report", required=True, type=Path)
    parser.add_argument("--line-numbers", action="store_true")
    args = parser.parse_args()

    template_path = resolve_template(args.template, Path.cwd())
    if template_path:
        print(f"Using template: {template_path}")
    else:
        print("Template not found; generating code docx without template.")
    try:
        build_docx(
            args.code_dir,
            args.output,
            template_path,
            args.line_numbers,
            args.software_name,
            args.software_version,
            args.source_manifest,
            args.originality_report,
        )
    except ValidationError as exc:
        print(f"Validation failed: {exc}", file=sys.stderr)
        return 1
    print(f"Wrote {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
