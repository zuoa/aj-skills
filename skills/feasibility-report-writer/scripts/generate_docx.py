#!/usr/bin/env python3
"""Assemble the feasibility report Markdown into a styled DOCX.

Pipeline:
  1. parse optional front matter (cover/header) and **strip inline `[信源N]`
     citation markers from the body** — citations live only in the standalone
     `参考文献与信源清单.md`, never in the report text;
  2. (optional) render `参考文献与信源清单.md` from sources.json + run the
     citation audit (body must be clean; sources must be complete);
  3. convert report Markdown -> DOCX via pypandoc (reference docx for base
     structure);
  4. enforce the exact format of `可行性分析报告.doc` with python-docx: A4 page,
     cover page, 三号/小三/四号 chapter/body/caption sizes, 黑体/宋体 CJK fonts,
     首行缩进2字符, 页眉(报告全称) + 页脚「第 X 页」;
  5. write a sha256 sidecar so unchanged inputs are skipped on re-run.

Run from the user's working directory; outputs are relative to it. Scripts and
the venv are addressed via the skill directory (`{baseDir}`).

Usage:
  python3 {baseDir}/scripts/generate_docx.py \
      --input outputs/可行性分析报告_xxx_v1.0.md \
      --output outputs/可行性分析报告_xxx_v1.0.docx \
      --sources outputs/research/sources.json [--report outputs/quality_report.json]
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from pathlib import Path
from typing import Any

GENERATOR_VERSION = "3"
DEFAULT_REF_DOC = Path(__file__).resolve().parent.parent / "templates" / "research-report-reference.docx"
DEFAULT_HEADER = "宁波市公益类科技计划项目可行性分析报告"
CITE_RE = re.compile(r"\[信源\s*(\d+)\]")
PENDING_RE = re.compile(r"【待(?:确认|验证)[：:]?[^】]*】")
PENDING_COLOR = "0070C0"  # 待确认/待验证 标记的蓝色

# --- format map (extracted from 可行性分析报告.doc) -------------------------
# Chinese 字号 -> pt: 小一=24, 一号=26, 二号=22, 小二=18, 三号=16, 小三=15,
# 四号=14, 小四=12, 五号=10.5
COVER_TITLE_PT = 24      # 小一 — 封面项目名称 / 副标题
COVER_META_PT = 14       # 四号 — 封面承担单位 / 申报年月
CHAPTER_PT = 16          # 三号 — 章标题 一、…（限N字）
BODY_PT = 15             # 小三 — 正文 / 条 / 节标题
CAPTION_PT = 14          # 四号 — 图题 图1 …
HEADER_PT = 12           # 小四 — 页眉 / 页脚
INDENT_2CH_PT = 30       # ≈ 2 个汉字（15pt × 2），首行缩进
CJK_BODY = "宋体"
CJK_HEADING = "黑体"
LATIN = "Times New Roman"

CHAPTER_RE = re.compile(r"^\s*[一二三四五六]、")
SECTION_RE = re.compile(r"^\s*（[一二三四五六七八九十]+）")
FIGURE_RE = re.compile(r"^\s*图\s*\d+")
FRONT_MATTER_RE = re.compile(r"\A---\s*\n(.*?)\n---\s*\n?(.*)\Z", re.S)


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Generate styled DOCX from the assembled report Markdown")
    p.add_argument("--input", required=True, type=Path, help="Assembled report .md file")
    p.add_argument("--output", required=True, type=Path, help="Output .docx path")
    p.add_argument("--sources", type=Path, help="Optional sources.json for the citation list")
    p.add_argument("--reference-doc", type=Path, default=DEFAULT_REF_DOC, help="pandoc reference docx")
    p.add_argument("--report", type=Path, help="Optional quality_report.json to merge results into")
    p.add_argument("--overwrite", action="store_true", help="Regenerate even if sha256 matches")
    p.add_argument("--with-markdown", action="store_true", help="Copy cleaned md next to output (compat flag)")
    # cover / header overrides (front-matter in the .md takes precedence; these override it)
    p.add_argument("--title", help="Cover project name (overrides front-matter)")
    p.add_argument("--unit", help="Cover undertaking unit (overrides front-matter)")
    p.add_argument("--date", help="Cover submission year-month (overrides front-matter)")
    p.add_argument("--subtitle", help="Cover subtitle (default 项目可行性报告)")
    p.add_argument("--header-text", help=f"Running header text (default {DEFAULT_HEADER})")
    p.add_argument("--no-cover", action="store_true", help="Skip building a cover page")
    p.add_argument("--no-header", action="store_true", help="Skip the running header")
    p.add_argument("--no-footer", action="store_true", help="Skip the 第 X 页 page-number footer")
    return p.parse_args()


def file_sha256(path: Path) -> str:
    if not path.exists():
        return ""
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def stable_hash(md_text: str, ref_hash: str, sources_hash: str) -> str:
    payload = {
        "md_sha256": hashlib.sha256(md_text.encode("utf-8")).hexdigest(),
        "reference_doc_sha256": ref_hash,
        "sources_sha256": sources_hash,
        "generator_version": GENERATOR_VERSION,
    }
    s = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(s.encode("utf-8")).hexdigest()


def sidecar_path(output: Path) -> Path:
    return output.with_suffix(output.suffix + ".sha256")


def should_skip(output: Path, h: str) -> bool:
    sc = sidecar_path(output)
    if not output.exists() or not sc.exists():
        return False
    return sc.read_text(encoding="utf-8").strip() == h


def parse_front_matter(md: str) -> tuple[dict[str, str], str]:
    """Pull a simple leading ``---\\nkey: value\\n---`` block; return (meta, rest)."""
    m = FRONT_MATTER_RE.match(md)
    if not m:
        return {}, md
    meta: dict[str, str] = {}
    for line in m.group(1).splitlines():
        if ":" not in line:
            continue
        k, _, v = line.partition(":")
        v = v.strip().strip('"').strip("'").strip()
        if k.strip():
            meta[k.strip()] = v
    return meta, m.group(2)


def strip_citations(md: str) -> str:
    """Remove inline `[信源N]` markers (and any preceding space) from body text."""
    return re.sub(r"\s*\[信源\s*\d+\]", "", md)


def render_sources_md(sources: list[dict]) -> str:
    lines = [
        "# 参考文献与信源清单",
        "",
        "> 人工复核用：按下表逐条核对——其「引用要点」指向的报告章节与引用内容，"
        "是否与正文一致；点链接对照原文是否支撑该数据/结论。正文不出现引用标记，"
        "信源全部集中在本清单。",
        "",
        "| 序号 | 类型 | 标题 | 发布机构 | 日期(发布/访问) | 链接 | 引用要点（报告章节·引用内容） | 核验状态 |",
        "|---|---|---|---|---|---|---|---|",
    ]
    for s in sources:
        sid = s.get("id", "")
        stype = s.get("type", "")
        title = str(s.get("title", "")).replace("|", "\\|")
        issuer = str(s.get("issuer", "")).replace("|", "\\|")
        date = s.get("date", "")
        accessed = s.get("accessed", "")
        date_cell = f"{date} / 访问 {accessed}" if accessed else f"{date}"
        url = s.get("url", "")
        link = f"[{url}]({url})" if url else "—"
        used = str(s.get("used_in", "")).replace("|", "\\|")
        claim = str(s.get("claim", "")).replace("|", "\\|").replace("\n", " ")
        status = s.get("status", "")
        lines.append(f"| {sid} | {stype} | {title} | {issuer} | {date_cell} | {link} | {used}·{claim} | {status} |")
    lines.append("")
    return "\n".join(lines)


def citation_audit(md_text: str, sources: list[dict]) -> dict:
    """Body must carry no `[信源N]`; every source should have id+title+used_in."""
    remaining = sorted({int(n) for n in CITE_RE.findall(md_text)})
    missing_title = [s.get("id") for s in sources if not str(s.get("title", "")).strip()]
    missing_usage = [s.get("id") for s in sources if not str(s.get("used_in", "")).strip()]
    return {
        "policy": "正文不出现 [信源N]；信源统一在独立清单，靠『引用要点（章节·引用内容）』与正文对应",
        "body_citation_markers_remaining": remaining,
        "body_clean": not remaining,
        "sources_count": len(sources),
        "sources_missing_title": missing_title,
        "sources_missing_used_in": missing_usage,
        "pass": not remaining,
    }


# --- python-docx styling helpers ------------------------------------------
def _style_run(run, size_pt: float, bold: bool, cjk: str = CJK_BODY, latin: str = LATIN) -> None:
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Pt  # type: ignore

    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)
    rFonts.set(qn("w:eastAsia"), cjk)


def _classify(para) -> str:
    name = (para.style.name or "") if para.style else ""
    text = (para.text or "").strip()
    if name == "Heading 1" or CHAPTER_RE.match(text):
        return "chapter"
    if name == "Heading 2" or SECTION_RE.match(text):
        return "section"
    if FIGURE_RE.match(text) and len(text) <= 30:
        return "figure"
    return "body"


def style_body(doc) -> None:
    """Apply the .doc style map to every paragraph (cover is added later, separately)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Pt  # type: ignore

    JUSTIFY, CENTER, LEFT = WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT
    for p in doc.paragraphs:
        kind = _classify(p)
        pf = p.paragraph_format
        if kind == "chapter":
            pf.alignment, pf.first_line_indent = LEFT, Pt(0)
            pf.space_before, pf.space_after, pf.line_spacing = Pt(8), Pt(4), 1.15
            size, bold, cjk = CHAPTER_PT, False, CJK_HEADING
        elif kind == "section":
            pf.alignment, pf.first_line_indent = JUSTIFY, Pt(INDENT_2CH_PT)
            pf.space_before, pf.space_after, pf.line_spacing = Pt(4), Pt(2), 1.15
            size, bold, cjk = BODY_PT, True, CJK_BODY
        elif kind == "figure":
            pf.alignment, pf.first_line_indent = CENTER, Pt(0)
            pf.space_before, pf.space_after = Pt(2), Pt(6)
            size, bold, cjk = CAPTION_PT, False, CJK_BODY
        else:
            pf.alignment, pf.first_line_indent = JUSTIFY, Pt(INDENT_2CH_PT)
            pf.space_before, pf.space_after, pf.line_spacing = Pt(0), Pt(0), 1.15
            size, bold, cjk = BODY_PT, False, CJK_BODY
        for r in p.runs:
            _style_run(r, size, bold, cjk)


def _set_element_text(r_elem, text: str) -> None:
    """Replace all <w:t> in a run element with a single text node."""
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    for child in list(r_elem):
        if child.tag == qn("w:t"):
            r_elem.remove(child)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r_elem.append(t)


def _mark_pending(r_elem) -> None:
    """Add italic + blue color to a run element (for 待确认/待验证 markers)."""
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    rPr = r_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r_elem.insert(0, rPr)
    if rPr.find(qn("w:i")) is None:
        rPr.append(OxmlElement("w:i"))
    color = rPr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rPr.append(color)
    color.set(qn("w:val"), PENDING_COLOR)


def highlight_pending(doc) -> None:
    """Render inline 【待确认：…】/【待验证：…】 markers as blue italic.

    Works at the paragraph level on the concatenated run text so markers that
    pandoc split across several runs are still caught; each affected run is split
    at the marker boundaries and the marker portion gains italic + blue.
    """
    from copy import deepcopy
    from docx.oxml.ns import qn  # type: ignore

    for p in doc.paragraphs:
        runs = list(p._element.findall(qn("w:r")))
        run_text = ["".join((t.text or "") for t in r.findall(qn("w:t"))) for r in runs]
        full = "".join(run_text)
        matches = list(PENDING_RE.finditer(full))
        if not matches:
            continue
        marker_pos = [False] * len(full)
        for m in matches:
            for k in range(m.start(), m.end()):
                marker_pos[k] = True
        pos = 0
        for r, text in zip(runs, run_text):
            if not text:
                continue
            chunks: list[tuple[str, bool]] = []
            i = 0
            while i < len(text):
                cur = marker_pos[pos + i]
                j = i + 1
                while j < len(text) and marker_pos[pos + j] == cur:
                    j += 1
                chunks.append((text[i:j], cur))
                i = j
            pos += len(text)
            if len(chunks) == 1 and not chunks[0][1]:
                continue  # plain run, untouched
            for ctext, is_marker in chunks:
                nr = deepcopy(r)
                _set_element_text(nr, ctext)
                if is_marker:
                    _mark_pending(nr)
                r.addprevious(nr)
            r.getparent().remove(r)


def _insert_cover_para(doc, before_elem, text, size_pt, bold, cjk, align, *, empty=False, space_after=0.0) -> None:
    from docx.shared import Pt  # type: ignore

    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.space_before, pf.space_after, pf.line_spacing = Pt(0), Pt(space_after), 1.15
    if not empty and text:
        _style_run(p.add_run(text), size_pt, bold, cjk)
    before_elem.addprevious(p._element)


def build_cover(doc, cover: dict[str, str], first_body_para) -> None:
    """Prepend a centered cover page; push the body onto the next page."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

    center = WD_ALIGN_PARAGRAPH.CENTER
    title = cover.get("title", "").strip()
    subtitle = cover.get("subtitle", "项目可行性报告").strip()
    unit = cover.get("unit", "").strip()
    date = cover.get("date", "").strip()
    if not (title or unit or date):
        return
    before = first_body_para._element
    for _ in range(2):
        _insert_cover_para(doc, before, "", 0, False, CJK_BODY, center, empty=True)
    if title:
        _insert_cover_para(doc, before, title, COVER_TITLE_PT, True, CJK_HEADING, center, space_after=6)
    _insert_cover_para(doc, before, subtitle, COVER_TITLE_PT, False, CJK_HEADING, center, space_after=6)
    for _ in range(6):
        _insert_cover_para(doc, before, "", 0, False, CJK_BODY, center, empty=True)
    if unit:
        _insert_cover_para(doc, before, unit, COVER_META_PT, True, CJK_BODY, center, space_after=4)
    if date:
        _insert_cover_para(doc, before, date, COVER_META_PT, True, CJK_BODY, center)
    first_body_para.paragraph_format.page_break_before = True


def _add_page_field(paragraph) -> None:
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore

    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(i); run._r.append(e)


def setup_page(doc, header_text: str, *, no_header: bool, no_footer: bool) -> None:
    """A4 + margins; cover (first page) clean; body pages carry header + 第 X 页 footer."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.shared import Mm  # type: ignore

    center = WD_ALIGN_PARAGRAPH.CENTER
    for section in doc.sections:
        section.page_width, section.page_height = Mm(210), Mm(297)
        section.top_margin, section.bottom_margin = Mm(25.4), Mm(25.4)
        section.left_margin, section.right_margin = Mm(31.8), Mm(31.8)
        section.different_first_page_header_footer = True  # cover (page 1) stays clean

        if not no_header and header_text:
            hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
            hp.text = ""
            hp.alignment = center
            _style_run(hp.add_run(header_text), HEADER_PT, True, CJK_BODY)

        if not no_footer:
            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            fp.text = ""
            fp.alignment = center
            _style_run(fp.add_run("第 "), HEADER_PT, False, CJK_BODY)
            _add_page_field(fp)
            _style_run(fp.add_run(" 页"), HEADER_PT, False, CJK_BODY)


def convert(md_text: str, output: Path, ref_doc: Path) -> None:
    import pypandoc  # type: ignore

    output.parent.mkdir(parents=True, exist_ok=True)
    extra_args = ["--reference-doc", str(ref_doc)] if ref_doc.exists() else []
    pypandoc.convert_text(
        md_text, to="docx", format="markdown",
        outputfile=str(output), extra_args=extra_args,
    )


def main() -> int:
    args = parse_args()
    if not args.input.is_file():
        print(f"ERROR: input not found: {args.input}", file=sys.stderr)
        return 2

    raw = args.input.read_text(encoding="utf-8")
    meta, body_md = parse_front_matter(raw)
    body_md = strip_citations(body_md)  # citations live only in the source list

    cover = {
        "title": args.title or meta.get("title") or meta.get("project_name", ""),
        "subtitle": args.subtitle or meta.get("subtitle", "项目可行性报告"),
        "unit": args.unit or meta.get("unit", ""),
        "date": args.date or meta.get("date", ""),
    }
    header_text = args.header_text or meta.get("header", DEFAULT_HEADER)

    ref_hash = file_sha256(args.reference_doc)
    sources: list[dict] = []
    sources_hash = ""
    sources_md_path: Path | None = None
    audit: dict[str, Any] | None = None

    if args.sources:
        if args.sources.is_file():
            try:
                sources = json.loads(args.sources.read_text(encoding="utf-8"))
            except json.JSONDecodeError as exc:
                print(f"ERROR: sources.json invalid: {exc}", file=sys.stderr)
                return 2
            sources_hash = file_sha256(args.sources)
            sources_md_path = args.sources.parent.parent / "参考文献与信源清单.md"
            sources_md_path.write_text(render_sources_md(sources), encoding="utf-8")
            print(f"Generated sources list: {sources_md_path}")
        else:
            print(f"WARNING: sources file not found, skipping citation list: {args.sources}")

    if sources:
        audit = citation_audit(body_md, sources)
        if audit["body_citation_markers_remaining"]:
            print(f"WARNING: 正文仍残留 [信源N] 标记（已自动剥离导出，但请回改源稿）：{audit['body_citation_markers_remaining']}")
        if audit["sources_missing_title"]:
            print(f"WARNING: 信源清单有条目缺标题：{audit['sources_missing_title']}")
        if audit["sources_missing_used_in"]:
            print(f"WARNING: 信源清单有条目缺『引用要点(章节)』：{audit['sources_missing_used_in']}")

    out_hash = stable_hash(body_md, ref_hash, sources_hash)
    if not args.overwrite and should_skip(args.output, out_hash):
        print(f"Skip (unchanged): {args.output}")
        return 0
    if args.output.exists() and not args.overwrite:
        print(f"Output exists with different hash; use --overwrite or bump version: {args.output}")
        return 2

    try:
        convert(body_md, args.output, args.reference_doc)
    except Exception as exc:  # noqa: BLE001
        print(f"ERROR: docx conversion failed: {exc}", file=sys.stderr)
        print("Install pandoc + pypandoc, or run `bash {baseDir}/scripts/setup_env.sh`.", file=sys.stderr)
        return 3

    try:
        from docx import Document  # type: ignore

        doc = Document(str(args.output))
        if doc.paragraphs:
            style_body(doc)
            highlight_pending(doc)
            if not args.no_cover:
                build_cover(doc, cover, doc.paragraphs[0])
        setup_page(doc, header_text, no_header=args.no_header, no_footer=args.no_footer)
        doc.save(str(args.output))
    except Exception as exc:  # noqa: BLE001
        print(f"WARNING: docx styling failed (output left unstyled): {exc}", file=sys.stderr)

    sidecar_path(args.output).write_text(out_hash + "\n", encoding="utf-8")

    if args.with_markdown:
        md_copy = args.output.with_suffix(".md")
        md_copy.write_text(body_md, encoding="utf-8")

    if args.report:
        existing = {}
        if args.report.exists():
            try:
                existing = json.loads(args.report.read_text(encoding="utf-8"))
            except (OSError, json.JSONDecodeError):
                existing = {}
        existing["docx"] = {"output": str(args.output), "sha256": out_hash,
                            "generator_version": GENERATOR_VERSION}
        if audit is not None:
            existing["citation_audit"] = audit
        if sources_md_path is not None:
            existing["sources_list"] = str(sources_md_path)
        args.report.parent.mkdir(parents=True, exist_ok=True)
        args.report.write_text(json.dumps(existing, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    print(f"Generated: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
