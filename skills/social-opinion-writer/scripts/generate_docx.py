#!/usr/bin/env python3
"""Generate a mainland-official-style docx from a markdown draft."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from normalize_cn_punctuation import normalize_markdown_text

TITLE_FONT = "方正小标宋简体"
BODY_FONT = "仿宋_GB2312"
LEVEL1_FONT = "黑体"
LEVEL2_FONT = "楷体_GB2312"

TITLE_SIZE = Pt(22)
BODY_SIZE = Pt(16)
REF_SIZE = Pt(14)
LINE_SPACING = Pt(28)
FIRST_LINE_INDENT = Pt(32)

CITATION_MARK = re.compile(r"\[\d+\]")


def set_run_font(run, east_asia_font: str, size, *, bold: bool = False) -> None:
    run.font.name = east_asia_font
    run.font.size = size
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def style_body_paragraph(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = LINE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = FIRST_LINE_INDENT
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def style_heading_paragraph(paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = LINE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_title(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = LINE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    run = paragraph.add_run(title.strip())
    set_run_font(run, TITLE_FONT, TITLE_SIZE)

    blank = doc.add_paragraph()
    fmt = blank.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = LINE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def classify_paragraph(text: str) -> str:
    stripped = text.strip()
    if re.match(r"^[一二三四五六七八九十]+、", stripped):
        return "level1"
    if re.match(r"^（[一二三四五六七八九十]+）", stripped):
        return "level2"
    if re.match(r"^\d+[.．、]", stripped):
        return "level3"
    if re.match(r"^【.+】$", stripped):
        return "section"
    return "body"


def is_structural_line(text: str) -> bool:
    return classify_paragraph(text) != "body"


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    kind = classify_paragraph(text)
    content = text.strip()

    if kind == "level1":
        style_heading_paragraph(paragraph)
        run = paragraph.add_run(content)
        set_run_font(run, LEVEL1_FONT, BODY_SIZE, bold=False)
        return

    if kind == "level2":
        style_heading_paragraph(paragraph)
        run = paragraph.add_run(content)
        set_run_font(run, LEVEL2_FONT, BODY_SIZE, bold=False)
        return

    if kind == "level3":
        style_heading_paragraph(paragraph)
        run = paragraph.add_run(content)
        set_run_font(run, BODY_FONT, BODY_SIZE, bold=True)
        return

    if kind == "section":
        style_heading_paragraph(paragraph)
        run = paragraph.add_run(content)
        set_run_font(run, LEVEL1_FONT, BODY_SIZE, bold=False)
        return

    style_body_paragraph(paragraph)
    # Render [n] citation marks as real superscript runs, like journal corner marks.
    pos = 0
    for match in CITATION_MARK.finditer(content):
        if match.start() > pos:
            run = paragraph.add_run(content[pos : match.start()])
            set_run_font(run, BODY_FONT, BODY_SIZE, bold=False)
        run = paragraph.add_run(match.group())
        set_run_font(run, BODY_FONT, BODY_SIZE, bold=False)
        run.font.superscript = True
        pos = match.end()
    if pos < len(content):
        run = paragraph.add_run(content[pos:])
        set_run_font(run, BODY_FONT, BODY_SIZE, bold=False)


def add_references_section(doc: Document, references_text: str) -> None:
    """Append the citation table as an appendix section, marks kept non-superscript."""

    blank = doc.add_paragraph()
    fmt = blank.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = LINE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)

    heading = doc.add_paragraph()
    style_heading_paragraph(heading)
    run = heading.add_run("附：引用来源")
    set_run_font(run, LEVEL1_FONT, BODY_SIZE, bold=False)

    for raw_line in references_text.splitlines():
        stripped = raw_line.strip()
        if not stripped or stripped.startswith("#"):
            continue
        if stripped.startswith("- "):
            stripped = stripped[2:].strip()
        paragraph = doc.add_paragraph()
        style_body_paragraph(paragraph)
        run = paragraph.add_run(stripped)
        set_run_font(run, BODY_FONT, REF_SIZE, bold=False)


def markdown_to_blocks(text: str) -> tuple[str, list[str]]:
    lines = text.splitlines()
    title = ""
    blocks: list[str] = []
    buffer: list[str] = []
    in_code_block = False

    def flush_buffer() -> None:
        nonlocal buffer
        if buffer:
            block = " ".join(part.strip() for part in buffer if part.strip()).strip()
            if block:
                blocks.append(block)
            buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue

        stripped = line.strip()
        if not stripped:
            flush_buffer()
            continue

        if not title:
            if stripped.startswith("# "):
                title = stripped[2:].strip()
                continue
            title = stripped.lstrip("#").strip()
            continue

        normalized = stripped
        if normalized.startswith("## "):
            normalized = normalized[3:].strip()
        elif normalized.startswith("### "):
            normalized = normalized[4:].strip()
        elif normalized.startswith("- "):
            normalized = normalized[2:].strip()

        if is_structural_line(normalized):
            flush_buffer()
            blocks.append(normalized)
            continue

        buffer.append(normalized)

    flush_buffer()
    return title, blocks


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, help="Path to final markdown file")
    parser.add_argument("--output", required=True, help="Path to output docx file")
    parser.add_argument("--references", help="Path to citation table markdown (11-references.md)")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    text = normalize_markdown_text(input_path.read_text(encoding="utf-8"))
    title, blocks = markdown_to_blocks(text)
    if not title:
        raise SystemExit("No title found in markdown input.")

    references_text = ""
    if args.references:
        references_path = Path(args.references)
        if references_path.is_file():
            references_text = normalize_markdown_text(references_path.read_text(encoding="utf-8"))
        else:
            print(f"References file not found, skipped: {references_path}")

    doc = Document()
    configure_document(doc)
    add_title(doc, title)
    for block in blocks:
        add_paragraph(doc, block)
    if references_text.strip():
        add_references_section(doc, references_text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Generated {output_path}")
    print(
        "Formatting basis: A4, 37/35/28/26mm margins, title 2nd-size Songti-style, "
        "body 3rd-size FangSong, exact 28pt line spacing."
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
